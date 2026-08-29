"""文档解析 / 生成工具：python-docx 与 docxtpl。"""

import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docxtpl import DocxTemplate

# 论文标题字号（黑体加粗）
_TITLE_FONT = "黑体"
_TITLE_SIZE_PT = 22  # 二号
# 一级=黑体三号(16pt)加粗；二级=黑体四号(14pt)加粗；三级=黑体小四(12pt)加粗
_HEADING_SIZE_PT = {1: 16, 2: 14, 3: 12}
# 正文=宋体小四(12pt)，1.5 倍行距，首行缩进两个字符（12pt × 2 = 24pt）
_BODY_FONT = "宋体"
_BODY_SIZE_PT = 12
_BODY_LINE_SPACING = 1.5
_BODY_FIRST_LINE_INDENT_PT = 24
# 参考文献正文=宋体五号(10.5pt)
_REFERENCE_SIZE_PT = 10.5
_MD_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")


def extract_text_from_docx(content: bytes) -> str:
    """解析 .docx 文件并提取纯文本。"""
    document = Document(BytesIO(content))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def render_docx_template(template_path: str, context: dict, output_path: str) -> None:
    """使用 docxtpl 渲染 Word 模板。

    Args:
        template_path: 带 Jinja2 占位符的 .docx 模板路径。
        context: 渲染上下文（论文标题、章节、图表等）。
        output_path: 输出文件路径。
    """
    template = DocxTemplate(template_path)
    template.render(context)
    template.save(output_path)


def build_paper_docx(title: str, sections: dict[str, str], output_path: str) -> None:
    """根据论文标题与章节正文生成 .docx 文档。

    章节正文为 Markdown 文本（WriterAgent 输出），生成时按标题层级套用字体：
    - 一级标题（##）：黑体三号加粗，且每个一级标题另起一页；
    - 二级标题（###）：黑体四号加粗；
    - 三级标题（#### 及更深）：黑体小四加粗；
    - 正文：宋体小四、1.5 倍行距、首行缩进两个字符；
    - 参考文献正文：宋体五号。
    """
    document = Document()
    _add_title(document, title)
    for section_title, body in sections.items():
        _append_section(document, section_title, body)
    document.save(output_path)


def _add_title(document: Document, title: str) -> None:
    """添加论文标题（黑体二号加粗，居中）。"""
    paragraph = document.add_heading(title, level=0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        _set_run_font(run, _TITLE_FONT, _TITLE_SIZE_PT)


def _append_section(document: Document, section_title: str, body: str) -> None:
    """按 Markdown 层级渲染单个章节：一级标题 + 各级标题 + 正文段落。"""
    lines = (body or "").splitlines()
    has_level1 = any(_markdown_level(line) == 1 for line in lines if line.strip())
    # 正文未给出章节标题时，回退使用数据库中的章节名作为一级标题
    if not has_level1:
        _add_heading(document, section_title, 1)
    is_references = _is_references_title(section_title)

    for line in lines:
        stripped = line.strip()
        if not stripped or stripped in ("---", "***", "___"):
            continue
        md_level = _markdown_level(stripped)
        if md_level:
            text = _MD_HEADING.match(stripped).group(2).strip()
            if not text:
                continue
            _add_heading(document, text, md_level)
            if md_level == 1:
                is_references = _is_references_title(text)
        elif is_references:
            _add_reference_paragraph(document, _clean_inline_markdown(stripped))
        else:
            _add_body_paragraph(document, _clean_inline_markdown(stripped))


def _markdown_level(line: str) -> int:
    """把 Markdown 标题行映射为论文层级：## -> 1；### -> 2；####+ -> 3；非标题返回 0。"""
    match = _MD_HEADING.match(line.strip())
    if not match:
        return 0
    markdown_level = len(match.group(1))
    if markdown_level <= 2:
        return 1
    if markdown_level == 3:
        return 2
    return 3


def _add_heading(document: Document, text: str, level: int) -> None:
    """添加指定层级标题，并套用黑体字号加粗格式；一级标题另起一页。"""
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        _set_run_font(run, _TITLE_FONT, _HEADING_SIZE_PT[level])
    if level == 1:
        paragraph.paragraph_format.page_break_before = True


def _add_body_paragraph(document: Document, text: str) -> None:
    """添加正文段落：宋体小四、1.5 倍行距、首行缩进两个字符。"""
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    _set_run_font(run, _BODY_FONT, _BODY_SIZE_PT, bold=False)
    paragraph.paragraph_format.line_spacing = _BODY_LINE_SPACING
    paragraph.paragraph_format.first_line_indent = Pt(_BODY_FIRST_LINE_INDENT_PT)


def _add_reference_paragraph(document: Document, text: str) -> None:
    """添加参考文献条目：宋体五号、1.5 倍行距、无首行缩进。"""
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    _set_run_font(run, _BODY_FONT, _REFERENCE_SIZE_PT, bold=False)
    paragraph.paragraph_format.line_spacing = _BODY_LINE_SPACING


def _is_references_title(title: str) -> bool:
    """判断章节标题是否属于参考文献章节。"""
    return "文献" in (title or "")


def _set_run_font(run, font_name: str, size_pt: int, bold: bool = True) -> None:
    """设置 run 的中英文字体、字号、加粗与黑色。"""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    # 中文字体需同时设置 eastAsia，否则中文仍可能回退默认字体
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _clean_inline_markdown(text: str) -> str:
    """去除正文中的加粗 / 斜体 / 行内代码标记，避免 Word 中显示原始符号。"""
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"_([^_]+)_", r"\1", text)
    return text
